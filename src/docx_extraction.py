
import docx
import os

from locations import base_location


doc = docx.Document(os.path.join(base_location, '/data/template_docx.docx'))
tables = doc.tables
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)
