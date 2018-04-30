# # with open('/Users/EmRo/Documents/LFAS/base_template.txt', 'rb') as f:
# #     for chunk in iter(lambda: f.read(4096), b''):
# #         txt_raw = chunk.decode('ascii', errors='ignore')
# #         txt_raw = txt_raw.replace('\r', ' ').replace('\n', ' ').replace('\t', ' ')
# #         print('---------------------------------')
# #         print(chunk.decode())
# #         print('---------------------------------')
# import re
#
# match_table = {
#     'total_fac_a': r'"Total Facility A Commitments".*\[(.*)\]',
#     'total_fac_b': r'"Total Facility B Commitments".*\[(.*)\]',
#     'total_fac_c': r'"Total Facility C Commitments".*\[(.*)\]',
#     'total_fac_d': r'"Total Facility D Commitments".*\[(.*)\]',
# }
#
#
# with open('/Users/EmRo/Documents/LFAS/base_template.txt', 'rb') as f:
#     raw_txt = f.read().decode('ascii', errors='ignore')
#
# matches = {}
# for var_name, src_pat in match_table.items():
#     mch = re.search(src_pat, raw_txt)
#     if mch:
#         print(mch.groups(1))
#     matches[var_name] = mch
#
# print(matches)
#
from io import BytesIO

from docx import Document

with open('/Users/EmRo/Documents/LFAS/template_docx.docx', 'rb') as f:
    source_stream = BytesIO(f.read())
document = Document(source_stream)
source_stream.close()


#
# fullText = []
# for para in document.paragraphs:
#     fullText.append(para.text)
#     print('\n'.join(fullText))
#     breack

def iter_headings(paragraphs):
    for paragraph in paragraphs:
        stl = paragraph.style.name
        if stl.startswith('Heading'):
            yield paragraph

counter = [0] * 10
for heading in iter_headings(document.paragraphs):
    lvl = int(heading.style.name[-1]) - 1
    counter[lvl] += 1
    counter[lvl+1:-1] = [0]*len(counter[lvl+1:-1])


